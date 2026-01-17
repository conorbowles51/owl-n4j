"""
Generate professional DOCX report from cost_analysis_report.txt file.

Creates a well-formatted Word document with proper sections and formatting.
"""

import sys
import re
from pathlib import Path
from datetime import datetime

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Owl brand colors
OWL_BLUE_800 = RGBColor(29, 77, 118)  # #1d4d76 - Primary dark blue
OWL_BLUE_900 = RGBColor(15, 47, 74)   # #0f2f4a - Darker blue for titles
OWL_BLUE_700 = RGBColor(36, 94, 143)  # #245e8f - Medium blue
OWL_BLUE_200 = RGBColor(158, 191, 223) # #9ebfdf - Light blue
OWL_PURPLE_500 = RGBColor(147, 51, 234)  # #9333ea - Accent purple
OWL_ORANGE_500 = RGBColor(249, 115, 22)  # #f97316 - Accent orange
OWL_GRAY_500 = RGBColor(107, 114, 128)   # #6b7280 - Medium gray
OWL_GRAY_400 = RGBColor(156, 163, 175)   # #9ca3af - Light gray


def parse_text_report(txt_path: Path):
    """Parse the text report and extract all data."""
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    data = {
        'documents_per_month': 100,
        'avg_chunks_per_document': 10.0,
        'avg_pages_per_document': 3.0,
        'generated_date': None,
        'ollama': {},
        'openai_models': [],
    }
    
    # Extract generated date
    date_match = re.search(r'Generated: (.+)', content)
    if date_match:
        data['generated_date'] = date_match.group(1).strip()
    
    # Extract documents per month
    docs_match = re.search(r'Processing (\d+) documents/month', content)
    if docs_match:
        data['documents_per_month'] = int(docs_match.group(1))
    
    # Extract chunks per document
    chunks_match = re.search(r'Average Chunks per Document: ([\d.]+)', content)
    if chunks_match:
        data['avg_chunks_per_document'] = float(chunks_match.group(1))
    
    # Extract Ollama data
    ollama_section = re.search(r'## SCENARIO 1: GPU SERVER.*?TOTAL MONTHLY COST: \$([\d.]+)', content, re.DOTALL)
    if ollama_section:
        time_match = re.search(r'Processing Time: ([\d.]+) seconds', ollama_section.group(0))
        entities_match = re.search(r'Entities Extracted: (\d+)', ollama_section.group(0))
        rel_match = re.search(r'Relationships Extracted: (\d+)', ollama_section.group(0))
        cost_doc_match = re.search(r'Cost per Document: \$([\d.]+)', ollama_section.group(0))
        
        data['ollama'] = {
            'time_per_chunk_seconds': float(time_match.group(1)) if time_match else 45.0,
            'entities_per_chunk': int(entities_match.group(1)) if entities_match else 8,
            'relationships_per_chunk': int(rel_match.group(1)) if rel_match else 6,
            'monthly_cost': float(ollama_section.group(1)),
            'cost_per_document': float(cost_doc_match.group(1)) if cost_doc_match else 14.40,
        }
    
    # Extract OpenAI models - updated to capture API costs and total monthly costs
    model_pattern = r'### Model: (.+?)\n.*?Test Results \(1 chunk\):.*?Processing Time: ([\d.]+) seconds.*?Cost per Chunk: \$([\d.]+).*?Token Usage: ([\d,]+) tokens.*?Entities Extracted: (\d+).*?Relationships Extracted: (\d+).*?Monthly Cost Projection.*?API Costs: \$([\d.]+).*?TOTAL MONTHLY COST: \$([\d.]+).*?Cost per Document: \$([\d.]+)'
    
    for match in re.finditer(model_pattern, content, re.DOTALL):
        model = {
            'name': match.group(1).strip(),
            'time_per_chunk_seconds': float(match.group(2)),
            'cost_per_chunk': float(match.group(3)),
            'tokens': int(match.group(4).replace(',', '')),
            'entities_per_chunk': int(match.group(5)),
            'relationships_per_chunk': int(match.group(6)),
            'api_cost_monthly': float(match.group(7)),
            'total_monthly_cost': float(match.group(8)),
            'cost_per_document': float(match.group(9)),
        }
        data['openai_models'].append(model)
    
    # Extract comparison data
    comparison_match = re.search(
        r'GPU Server \(Ollama\): \$([\d.]+)/month.*?Non-GPU Server \(([^)]+)\): \$([\d.]+)/month.*?Difference: \$([\d.]+)/month',
        content,
        re.DOTALL
    )
    if comparison_match:
        data['comparison'] = {
            'gpu_monthly': float(comparison_match.group(1)),
            'best_model': comparison_match.group(2).strip(),
            'openai_monthly': float(comparison_match.group(3)),
            'difference': float(comparison_match.group(4)),
        }
    
    return data


def calculate_metrics(data, model_data):
    """Calculate per-chunk, per-page, per-document metrics."""
    chunks = data['avg_chunks_per_document']
    pages = data['avg_pages_per_document']
    
    cost_per_chunk = model_data.get('cost_per_chunk', 0)
    time_per_chunk = model_data.get('time_per_chunk_seconds', 0)
    entities_per_chunk = model_data.get('entities_per_chunk', 0)
    relationships_per_chunk = model_data.get('relationships_per_chunk', 0)
    
    return {
        'per_chunk': {
            'cost': cost_per_chunk,
            'time_seconds': time_per_chunk,
            'entities': entities_per_chunk,
            'relationships': relationships_per_chunk,
        },
        'per_page': {
            'cost': cost_per_chunk * chunks / pages,
            'time_seconds': time_per_chunk * chunks / pages,
            'entities': entities_per_chunk * chunks / pages,
            'relationships': relationships_per_chunk * chunks / pages,
        },
        'per_document': {
            'cost': cost_per_chunk * chunks,
            'time_seconds': time_per_chunk * chunks,
            'entities': entities_per_chunk * chunks,
            'relationships': relationships_per_chunk * chunks,
        },
    }


def fmt_time(seconds):
    """Format time nicely."""
    if seconds < 60:
        return f"{seconds:.1f}s"
    elif seconds < 3600:
        return f"{seconds/60:.1f}min"
    else:
        return f"{seconds/3600:.2f}hrs"


def get_model_description(model_name: str, model_data: dict) -> tuple[str, str]:
    """
    Get description and behavior explanation for a model.
    Returns: (description, behavior_explanation)
    """
    entities = model_data.get('entities_per_chunk', 0)
    relationships = model_data.get('relationships_per_chunk', 0)
    cost_per_chunk = model_data.get('cost_per_chunk', 0)
    time_per_chunk = model_data.get('time_per_chunk_seconds', 0)
    tokens = model_data.get('tokens', 0)
    
    descriptions = {
        'gpt-4o-mini': (
            "A lightweight, cost-optimized version of GPT-4o designed for high-volume processing with minimal cost. "
            "Best suited for straightforward entity extraction tasks where speed and affordability are priorities.",
            f"This model processes quickly ({time_per_chunk:.1f}s) and costs only ${cost_per_chunk:.6f} per chunk, "
            f"making it the most economical option. However, its smaller architecture ({tokens:,} tokens used) results "
            f"in lower extraction depth ({entities} entities, {relationships} relationships per chunk). "
            f"The model prioritizes speed and cost over comprehensive analysis, which is why it extracts fewer entities "
            f"but still maintains good accuracy for basic entity recognition tasks."
        ),
        'gpt-4o': (
            "OpenAI's flagship multimodal model optimized for balanced performance across speed, cost, and quality. "
            "Designed as the successor to GPT-4 with improved efficiency and accuracy.",
            f"GPT-4o strikes an optimal balance: processing in {time_per_chunk:.1f}s with moderate cost "
            f"(${cost_per_chunk:.6f} per chunk) while extracting {entities} entities and {relationships} relationships. "
            f"The model uses {tokens:,} tokens, indicating efficient token utilization. Its architecture is optimized "
            f"for general-purpose tasks, providing consistent quality without excessive cost or latency. "
            f"This makes it ideal for production workloads where you need reliable performance without premium pricing."
        ),
        'gpt-5.2': (
            "OpenAI's most advanced model with maximum extraction capability and analytical depth. "
            "Designed for complex, high-stakes investigations requiring comprehensive entity and relationship mapping.",
            f"This model demonstrates the highest extraction capability ({entities} entities, {relationships} relationships) "
            f"but at a significant cost: ${cost_per_chunk:.6f} per chunk and {time_per_chunk:.1f}s processing time. "
            f"The model uses {tokens:,} tokens - nearly {tokens//1000}x more than smaller models - indicating deep analysis. "
            f"GPT-5.2's architecture prioritizes thoroughness over speed, performing extensive reasoning and pattern recognition. "
            f"This makes it ideal for complex fraud cases, intricate corporate structures, or when missing critical entities "
            f"has high consequences. The slower speed reflects the model's comprehensive multi-step reasoning process."
        ),
        'gpt-5-mini': (
            "An efficient entry-point to the GPT-5 series, offering superior extraction quality compared to GPT-4o-mini "
            "at a similarly low cost. Optimized for high-volume processing with better quality than its predecessor.",
            f"This model provides excellent value: {entities} entities and {relationships} relationships extracted "
            f"for only ${cost_per_chunk:.6f} per chunk. Processing takes {time_per_chunk:.1f}s, slower than GPT-4o-mini "
            f"but still fast. The model uses {tokens:,} tokens, suggesting improved reasoning capabilities over GPT-4o-mini "
            f"without the computational overhead of larger GPT-5 models. GPT-5-mini benefits from GPT-5's architectural "
            f"improvements (better instruction following, deeper reasoning) while maintaining efficiency, making it the "
            f"sweet spot for volume processing where you need quality improvements over GPT-4o-mini without premium costs."
        ),
        'gpt-5.1': (
            "A high-performance model in the GPT-5 series offering advanced extraction capabilities with strong "
            "relationship mapping. Designed for cases requiring detailed entity-relationship analysis without the "
            "maximum computational cost of GPT-5.2.",
            f"GPT-5.1 delivers strong extraction ({entities} entities, {relationships} relationships) at "
            f"${cost_per_chunk:.6f} per chunk with {time_per_chunk:.1f}s processing time. The model uses {tokens:,} tokens, "
            f"indicating substantial analytical depth. Notably, it extracts more relationships per entity than other models, "
            f"suggesting superior understanding of entity connections and contextual relationships. The model's architecture "
            f"balances the thoroughness of GPT-5.2 with better efficiency, making it ideal for cases requiring detailed "
            f"relationship mapping without the maximum extraction depth of GPT-5.2."
        ),
        'gpt-4.1': (
            "An enhanced version of GPT-4o with improved extraction capabilities and efficiency. "
            "Offers better quality than GPT-4o while maintaining competitive pricing and speed.",
            f"This model processes in {time_per_chunk:.1f}s with moderate cost (${cost_per_chunk:.6f} per chunk), "
            f"extracting {entities} entities and {relationships} relationships. Using {tokens:,} tokens, GPT-4.1 "
            f"demonstrates improved extraction depth over GPT-4o ({entities} vs 7 entities) while maintaining similar "
            f"speed and reasonable cost. The model benefits from refinements to GPT-4o's architecture, including better "
            f"instruction following and entity recognition. This makes it a strong middle-ground option when you need "
            f"quality improvements over GPT-4o without the higher costs of GPT-5 series models."
        ),
    }
    
    # Get description for model, or create generic one
    description, behavior = descriptions.get(model_name.lower(), (
        f"An OpenAI language model configured for entity and relationship extraction.",
        f"This model extracts {entities} entities and {relationships} relationships per chunk, "
        f"processing in {time_per_chunk:.1f}s at ${cost_per_chunk:.6f} per chunk. "
        f"The model uses {tokens:,} tokens, indicating {'high' if tokens > 10000 else 'moderate' if tokens > 5000 else 'efficient'} "
        f"computational requirements."
    ))
    
    return description, behavior


def fmt_currency(val):
    """Format currency."""
    return f"${val:,.4f}"


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def set_cell_shading(cell, color_hex):
    """Set table cell background color. color_hex should be like '1a365d' (no #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_logo_to_doc(doc, logo_path: Path, max_height: Inches = Inches(1.5)):
    """Add Owl logo to document, converting WEBP to PNG if needed."""
    try:
        if not logo_path.exists():
            return
        
        # Convert WEBP to PNG if needed (python-docx doesn't support WEBP)
        img = Image.open(logo_path)
        if img.format == 'WEBP':
            # Convert to PNG in memory
            png_buffer = io.BytesIO()
            if img.mode == 'RGBA':
                # Convert RGBA to RGB with white background
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3] if img.mode == 'RGBA' else None)
                rgb_img.save(png_buffer, format='PNG')
            else:
                img.save(png_buffer, format='PNG')
            png_buffer.seek(0)
            doc.add_picture(png_buffer, height=max_height)
        else:
            doc.add_picture(str(logo_path), height=max_height)
        
        # Center the logo
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Warning: Could not add logo: {e}")


def generate_docx(data, output_path: Path):
    """Generate DOCX report."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Cover Page
    # Add Owl logo at top
    logo_path = PROJECT_ROOT / 'frontend' / 'public' / 'owl-logo.webp'
    if logo_path.exists():
        add_logo_to_doc(doc, logo_path, max_height=Inches(1.5))
        doc.add_paragraph()  # Spacing after logo
        doc.add_paragraph()  # Extra spacing
    
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title = doc.add_heading('Infrastructure Cost Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = OWL_BLUE_900
    title_run.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph('Document Processing Solutions Comparison')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = OWL_BLUE_800
    
    doc.add_paragraph()  # Spacing
    
    para = doc.add_paragraph('Business Overview & Technical Performance Analysis')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    # Date
    date_para = doc.add_paragraph(f"Generated: {data.get('generated_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = OWL_GRAY_500
    
    para = doc.add_paragraph('Prepared by Owl Consultancy Group')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.size = Pt(11)
    para.runs[0].font.color.rgb = OWL_BLUE_800
    para.runs[0].bold = True
    
    add_page_break(doc)
    
    # SECTION 1: BUSINESS OVERVIEW
    heading1 = doc.add_heading('SECTION 1: BUSINESS OVERVIEW & RECOMMENDATIONS', 1)
    heading1.runs[0].font.color.rgb = OWL_BLUE_900
    
    # Executive Summary
    exec_summary = doc.add_heading('Executive Summary', 2)
    exec_summary.runs[0].font.color.rgb = OWL_BLUE_800
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(
        f"This comprehensive analysis evaluates two infrastructure solutions for processing "
        f"{data['documents_per_month']} documents per month (average {data['avg_pages_per_document']} pages per document). "
        f"The comparison assesses operational costs, processing speed, and extraction quality to provide "
        f"actionable recommendations for your document intelligence workflows."
    )
    summary_run.font.size = Pt(11)
    summary_para.paragraph_format.line_spacing = 1.6
    summary_para.paragraph_format.space_after = Pt(12)
    
    # Calculate cost ranges
    gpu_fixed_cost = data.get('comparison', {}).get('gpu_monthly', 1440)
    gpu_server_cost = gpu_fixed_cost  # GPU server: fixed cost only (no API fees)
    
    # Calculate OpenAI cost range
    openai_fixed_cost = 300  # $10/day * 30 days
    openai_monthly_costs = []
    openai_api_costs = []
    
    for model in data.get('openai_models', []):
        # Use total_monthly_cost if available (parsed from report), otherwise calculate
        if 'total_monthly_cost' in model:
            total_monthly = model['total_monthly_cost']
            api_cost = model.get('api_cost_monthly', total_monthly - openai_fixed_cost)
        else:
            # Fallback: calculate from cost_per_document (but this already includes server allocation)
            # Better to use cost_per_chunk * chunks_per_doc * docs_per_month + server
            if 'cost_per_chunk' in model:
                monthly_api_cost = model['cost_per_chunk'] * data['avg_chunks_per_document'] * data['documents_per_month']
                total_monthly = openai_fixed_cost + monthly_api_cost
                api_cost = monthly_api_cost
            else:
                # Last resort: use cost_per_document * docs_per_month (already includes server allocation)
                total_monthly = model.get('cost_per_document', 0) * data['documents_per_month']
                api_cost = total_monthly - openai_fixed_cost
        
        openai_monthly_costs.append(total_monthly)
        openai_api_costs.append(api_cost)
    
    if openai_monthly_costs:
        openai_min_cost = min(openai_monthly_costs)
        openai_max_cost = max(openai_monthly_costs)
        openai_min_api = min(openai_api_costs)
        openai_max_api = max(openai_api_costs)
        openai_cost_range = (openai_min_cost, openai_max_cost)
    else:
        # Fallback to comparison data if no model costs parsed
        openai_min_cost = data.get('comparison', {}).get('openai_monthly', 301)
        openai_max_cost = openai_min_cost
        openai_min_api = openai_min_cost - openai_fixed_cost
        openai_max_api = openai_min_api
        openai_cost_range = (openai_min_cost, openai_max_cost)
    
    # Determine winner and calculate preferred solution cost
    if data.get('comparison'):
        winner = "Non-GPU Server (OpenAI)" if openai_min_cost < gpu_server_cost else "GPU Server (Ollama)"
        savings = abs(gpu_server_cost - openai_min_cost) if openai_min_cost < gpu_server_cost else 0
    else:
        winner = "Non-GPU Server (OpenAI)" if openai_min_cost < gpu_server_cost else "GPU Server (Ollama)"
        savings = abs(gpu_server_cost - openai_min_cost) if openai_min_cost < gpu_server_cost else 0
    
    # Calculate preferred solution monthly cost
    if winner == "Non-GPU Server (OpenAI)":
        preferred_monthly_cost = openai_min_cost
    else:
        preferred_monthly_cost = gpu_server_cost
    
    # Case-based cost calculations (2-3 month case with 1000 documents ingested at start)
    case_documents = 1000
    case_months = 3
    case_duration_text = f"{case_months} months"
    
    # GPU Server case cost (runs for 3 months, documents processed at start but server runs continuously)
    gpu_case_cost = gpu_fixed_cost * case_months  # $1,440 * 3 = $4,320
    
    # Non-GPU Server case costs (server runs for 3 months, LLM processing done at start)
    openai_case_costs = []
    openai_case_api_costs = []
    
    for model in data.get('openai_models', []):
        # Case API cost = cost_per_chunk * chunks_per_doc * 1000 documents
        if 'cost_per_chunk' in model:
            case_api_cost = model['cost_per_chunk'] * data['avg_chunks_per_document'] * case_documents
        else:
            # Fallback: use cost_per_document * case_documents
            case_api_cost = model.get('cost_per_document', 0) * case_documents
        
        # Case server cost = $300/month * 3 months = $900
        case_server_cost = openai_fixed_cost * case_months
        case_total_cost = case_server_cost + case_api_cost
        
        openai_case_costs.append(case_total_cost)
        openai_case_api_costs.append(case_api_cost)
    
    if openai_case_costs:
        openai_case_min_cost = min(openai_case_costs)
        openai_case_max_cost = max(openai_case_costs)
        openai_case_min_api = min(openai_case_api_costs)
        openai_case_max_api = max(openai_case_api_costs)
    else:
        # Fallback: estimate from monthly costs
        case_server_cost = openai_fixed_cost * case_months
        openai_case_min_cost = case_server_cost + (openai_min_api * 10)  # 1000 docs / 100 docs per month = 10x
        openai_case_max_cost = case_server_cost + (openai_max_api * 10)
        openai_case_min_api = openai_min_api * 10
        openai_case_max_api = openai_max_api * 10
    
    # Calculate preferred solution case cost
    if winner == "Non-GPU Server (OpenAI)":
        preferred_case_cost = openai_case_min_cost
    else:
        preferred_case_cost = gpu_case_cost
    
    # Case savings
    case_savings = abs(gpu_case_cost - openai_case_min_cost) if openai_case_min_cost < gpu_case_cost else 0
    
    # Cost Ranges Section
    cost_label = doc.add_paragraph('SOLUTION COST RANGES:')
    cost_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cost_label.runs[0].font.size = Pt(14)
    cost_label.runs[0].font.color.rgb = OWL_BLUE_700
    cost_label.runs[0].bold = True
    
    # Case scenario note
    case_note = doc.add_paragraph(f'Note: Case scenario = {case_documents:,} documents over {case_duration_text}')
    case_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_note.runs[0].font.size = Pt(10)
    case_note.runs[0].font.color.rgb = OWL_GRAY_500
    case_note.runs[0].italic = True
    
    doc.add_paragraph()  # Spacing
    
    # GPU Server Cost
    gpu_label = doc.add_paragraph('GPU Server (Ollama):')
    gpu_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    gpu_label.runs[0].font.size = Pt(12)
    gpu_label.runs[0].font.color.rgb = OWL_BLUE_800
    gpu_label.runs[0].bold = True
    
    gpu_cost_para = doc.add_paragraph()
    gpu_cost_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    gpu_cost_para.add_run("Server: ").font.size = Pt(11)
    gpu_cost_para.add_run(fmt_currency(gpu_fixed_cost)).font.size = Pt(11)
    gpu_cost_para.add_run("/month (fixed) + ").font.size = Pt(11)
    gpu_cost_para.add_run("$0 LLM processing").font.size = Pt(11)
    gpu_cost_para.add_run(" = ").font.size = Pt(11)
    total_gpu_run = gpu_cost_para.add_run(fmt_currency(gpu_server_cost))
    total_gpu_run.font.size = Pt(13)
    total_gpu_run.font.color.rgb = OWL_BLUE_900
    total_gpu_run.bold = True
    gpu_cost_para.add_run("/month").font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Non-GPU Server Cost Range
    openai_label = doc.add_paragraph('Non-GPU Server (OpenAI):')
    openai_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    openai_label.runs[0].font.size = Pt(12)
    openai_label.runs[0].font.color.rgb = OWL_BLUE_800
    openai_label.runs[0].bold = True
    
    if openai_min_cost == openai_max_cost:
        # Single cost (only one model or all models same cost)
        openai_cost_para = doc.add_paragraph()
        openai_cost_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        openai_cost_para.add_run("Server: ").font.size = Pt(11)
        openai_cost_para.add_run(fmt_currency(openai_fixed_cost)).font.size = Pt(11)
        openai_cost_para.add_run("/month (fixed) + ").font.size = Pt(11)
        openai_cost_para.add_run(fmt_currency(openai_min_api)).font.size = Pt(11)
        openai_cost_para.add_run(" LLM processing").font.size = Pt(11)
        openai_cost_para.add_run(" = ").font.size = Pt(11)
        total_openai_run = openai_cost_para.add_run(fmt_currency(openai_min_cost))
        total_openai_run.font.size = Pt(13)
        total_openai_run.font.color.rgb = OWL_BLUE_900
        total_openai_run.bold = True
        openai_cost_para.add_run("/month").font.size = Pt(11)
    else:
        # Cost range
        openai_cost_para = doc.add_paragraph()
        openai_cost_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        openai_cost_para.add_run("Server: ").font.size = Pt(11)
        openai_cost_para.add_run(fmt_currency(openai_fixed_cost)).font.size = Pt(11)
        openai_cost_para.add_run("/month (fixed) + ").font.size = Pt(11)
        openai_cost_para.add_run("LLM processing: ").font.size = Pt(11)
        openai_cost_para.add_run(fmt_currency(openai_min_api)).font.size = Pt(11)
        openai_cost_para.add_run(" - ").font.size = Pt(11)
        openai_cost_para.add_run(fmt_currency(openai_max_api)).font.size = Pt(11)
        openai_cost_para.add_run("/month").font.size = Pt(11)
        openai_cost_para.add_run(" = ").font.size = Pt(11)
        total_range_run = openai_cost_para.add_run(fmt_currency(openai_min_cost))
        total_range_run.font.size = Pt(13)
        total_range_run.font.color.rgb = OWL_BLUE_900
        total_range_run.bold = True
        openai_cost_para.add_run(" - ").font.size = Pt(11)
        total_max_run = openai_cost_para.add_run(fmt_currency(openai_max_cost))
        total_max_run.font.size = Pt(13)
        total_max_run.font.color.rgb = OWL_BLUE_900
        total_max_run.bold = True
        openai_cost_para.add_run("/month").font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Recommendation
    rec_label = doc.add_paragraph('RECOMMENDED SOLUTION:')
    rec_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rec_label.runs[0].font.size = Pt(12)
    rec_label.runs[0].font.color.rgb = OWL_BLUE_700
    rec_label.runs[0].bold = True
    
    para = doc.add_paragraph(winner)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.size = Pt(20)
    para.runs[0].font.color.rgb = OWL_BLUE_900
    para.runs[0].bold = True
    
    # Monthly Cost and Savings - show range
    cost_label = doc.add_paragraph('MONTHLY COST:')
    cost_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cost_label.runs[0].font.size = Pt(11)
    cost_label.runs[0].font.color.rgb = OWL_BLUE_700
    cost_label.runs[0].bold = True
    
    # Show range: min (server + cheapest model) to max (server + most expensive model)
    if winner == "Non-GPU Server (OpenAI)":
        # Show range for Non-GPU server
        if openai_min_cost != openai_max_cost:
            monthly_cost_para = doc.add_paragraph()
            monthly_cost_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            monthly_cost_min_run = monthly_cost_para.add_run(fmt_currency(openai_min_cost))
            monthly_cost_min_run.font.size = Pt(16)
            monthly_cost_min_run.font.color.rgb = OWL_ORANGE_500
            monthly_cost_min_run.bold = True
            monthly_cost_para.add_run(" - ").font.size = Pt(16)
            monthly_cost_max_run = monthly_cost_para.add_run(fmt_currency(openai_max_cost))
            monthly_cost_max_run.font.size = Pt(16)
            monthly_cost_max_run.font.color.rgb = OWL_ORANGE_500
            monthly_cost_max_run.bold = True
        else:
            # Single cost
            para = doc.add_paragraph(fmt_currency(openai_min_cost))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].font.size = Pt(18)
            para.runs[0].font.color.rgb = OWL_ORANGE_500
            para.runs[0].bold = True
    else:
        # GPU server - single cost (no range)
        para = doc.add_paragraph(fmt_currency(gpu_server_cost))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(18)
        para.runs[0].font.color.rgb = OWL_ORANGE_500
        para.runs[0].bold = True
    
    if savings > 0:
        savings_label = doc.add_paragraph('ESTIMATED MONTHLY SAVINGS:')
        savings_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        savings_label.runs[0].font.size = Pt(11)
        savings_label.runs[0].font.color.rgb = OWL_BLUE_700
        savings_label.runs[0].bold = True
        
        para = doc.add_paragraph(fmt_currency(savings))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(18)
        para.runs[0].font.color.rgb = OWL_PURPLE_500  # Purple for savings - different from cost
        para.runs[0].bold = True
    
    # Case Cost and Savings
    doc.add_paragraph()  # Spacing
    
    case_cost_label = doc.add_paragraph('PER-CASE COST:')
    case_cost_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_cost_label.runs[0].font.size = Pt(11)
    case_cost_label.runs[0].font.color.rgb = OWL_BLUE_700
    case_cost_label.runs[0].bold = True
    
    # Show cost range if there's a range for Non-GPU
    if winner == "Non-GPU Server (OpenAI)" and openai_case_min_cost != openai_case_max_cost:
        case_cost_para = doc.add_paragraph()
        case_cost_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_cost_min_run = case_cost_para.add_run(fmt_currency(openai_case_min_cost))
        case_cost_min_run.font.size = Pt(16)
        case_cost_min_run.font.color.rgb = OWL_ORANGE_500
        case_cost_min_run.bold = True
        case_cost_para.add_run(" - ").font.size = Pt(16)
        case_cost_para.add_run(fmt_currency(openai_case_max_cost)).font.size = Pt(16)
        case_cost_para.runs[2].font.color.rgb = OWL_ORANGE_500
        case_cost_para.runs[2].bold = True
    else:
        case_cost_para = doc.add_paragraph(fmt_currency(preferred_case_cost))
        case_cost_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_cost_para.runs[0].font.size = Pt(18)
        case_cost_para.runs[0].font.color.rgb = OWL_ORANGE_500  # Orange for case cost
        case_cost_para.runs[0].bold = True
    
    if case_savings > 0:
        case_savings_label = doc.add_paragraph('ESTIMATED PER-CASE SAVINGS:')
        case_savings_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_savings_label.runs[0].font.size = Pt(11)
        case_savings_label.runs[0].font.color.rgb = OWL_BLUE_700
        case_savings_label.runs[0].bold = True
        
        para = doc.add_paragraph(fmt_currency(case_savings))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(18)
        para.runs[0].font.color.rgb = OWL_PURPLE_500  # Purple for case savings
        para.runs[0].bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Cost Comparison Table
    cost_heading = doc.add_heading('Cost Comparison Overview', 2)
    cost_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Solution'
    header_cells[1].text = 'Monthly Cost'
    header_cells[2].text = 'Cost per Document'
    header_cells[3].text = 'Processing Speed'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '1d4d76')  # owl-blue-800
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # GPU Server row
    row1 = table.rows[1].cells
    row1[0].text = 'GPU Server (Ollama)'
    row1[0].paragraphs[0].runs[0].font.bold = True
    row1[1].text = fmt_currency(data.get('comparison', {}).get('gpu_monthly', 1440))
    row1[2].text = fmt_currency(data['ollama'].get('cost_per_document', 14.40))
    row1[3].text = fmt_time(data['ollama'].get('time_per_chunk_seconds', 45) * data['avg_chunks_per_document'])
    
    # Non-GPU Server row
    row2 = table.rows[2].cells
    row2[0].text = 'Non-GPU Server (OpenAI)'
    row2[0].paragraphs[0].runs[0].font.bold = True
    row2[1].text = fmt_currency(data.get('comparison', {}).get('openai_monthly', 301))
    row2[2].text = fmt_currency(min([m.get('cost_per_document', 3) for m in data['openai_models']], default=3))
    row2[3].text = fmt_time(min([m.get('time_per_chunk_seconds', 30) for m in data['openai_models']], default=30) * data['avg_chunks_per_document'])
    
    # Highlight winner with light blue
    if data.get('comparison', {}).get('openai_monthly', 301) < data.get('comparison', {}).get('gpu_monthly', 1440):
        for cell in row2:
            set_cell_shading(cell, 'e8f0f7')  # owl-blue-50 (very light blue)
    
    doc.add_paragraph()  # Spacing
    
    # Recommendation
    rec_heading = doc.add_heading('Recommendation', 2)
    rec_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    para = doc.add_paragraph()
    para.add_run(f"For {data['documents_per_month']} documents per month, we recommend the ").font.size = Pt(11)
    run = para.add_run(f"{winner}")
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(" solution.").font.size = Pt(11)
    
    para = doc.add_paragraph()
    para.add_run("This solution provides the optimal balance of ").font.size = Pt(11)
    run = para.add_run("cost efficiency")
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(", ").font.size = Pt(11)
    run = para.add_run("processing speed")
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(", and ").font.size = Pt(11)
    run = para.add_run("extraction quality")
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(" for your current volume.").font.size = Pt(11)
    
    if savings > 0:
        para = doc.add_paragraph()
        para.add_run(f"Monthly cost: ").font.size = Pt(11)
        if winner == "Non-GPU Server (OpenAI)" and openai_min_cost != openai_max_cost:
            # Show range for Non-GPU server
            cost_min_run = para.add_run(fmt_currency(openai_min_cost))
            cost_min_run.font.size = Pt(11)
            cost_min_run.bold = True
            cost_min_run.font.color.rgb = OWL_ORANGE_500
            para.add_run(" - ").font.size = Pt(11)
            cost_max_run = para.add_run(fmt_currency(openai_max_cost))
            cost_max_run.font.size = Pt(11)
            cost_max_run.bold = True
            cost_max_run.font.color.rgb = OWL_ORANGE_500
        else:
            # Single cost
            cost_run = para.add_run(fmt_currency(preferred_monthly_cost))
            cost_run.font.size = Pt(11)
            cost_run.bold = True
            cost_run.font.color.rgb = OWL_ORANGE_500
        para.add_run(". Estimated monthly savings: ").font.size = Pt(11)
        savings_run = para.add_run(fmt_currency(savings))
        savings_run.font.size = Pt(11)
        savings_run.bold = True
        savings_run.font.color.rgb = OWL_PURPLE_500  # Purple for savings
        para.add_run(" compared to the alternative solution.").font.size = Pt(11)
    else:
        para = doc.add_paragraph()
        para.add_run(f"Monthly cost: ").font.size = Pt(11)
        if winner == "Non-GPU Server (OpenAI)" and openai_min_cost != openai_max_cost:
            # Show range for Non-GPU server
            cost_min_run = para.add_run(fmt_currency(openai_min_cost))
            cost_min_run.font.size = Pt(11)
            cost_min_run.bold = True
            cost_min_run.font.color.rgb = OWL_ORANGE_500
            para.add_run(" - ").font.size = Pt(11)
            cost_max_run = para.add_run(fmt_currency(openai_max_cost))
            cost_max_run.font.size = Pt(11)
            cost_max_run.bold = True
            cost_max_run.font.color.rgb = OWL_ORANGE_500
        else:
            # Single cost
            cost_run = para.add_run(fmt_currency(preferred_monthly_cost))
            cost_run.font.size = Pt(11)
            cost_run.bold = True
            cost_run.font.color.rgb = OWL_ORANGE_500
        para.add_run(".").font.size = Pt(11)
    
    # Case cost in recommendation text
    if case_savings > 0:
        para = doc.add_paragraph()
        if winner == "Non-GPU Server (OpenAI)" and openai_case_min_cost != openai_case_max_cost:
            para.add_run(f"Per-case cost ({case_duration_text}, {case_documents:,} documents): ").font.size = Pt(11)
            case_cost_run = para.add_run(fmt_currency(openai_case_min_cost))
            case_cost_run.font.size = Pt(11)
            case_cost_run.bold = True
            case_cost_run.font.color.rgb = OWL_ORANGE_500
            para.add_run(" - ").font.size = Pt(11)
            case_cost_max_run = para.add_run(fmt_currency(openai_case_max_cost))
            case_cost_max_run.font.size = Pt(11)
            case_cost_max_run.bold = True
            case_cost_max_run.font.color.rgb = OWL_ORANGE_500
        else:
            para.add_run(f"Per-case cost ({case_duration_text}, {case_documents:,} documents): ").font.size = Pt(11)
            case_cost_run = para.add_run(fmt_currency(preferred_case_cost))
            case_cost_run.font.size = Pt(11)
            case_cost_run.bold = True
            case_cost_run.font.color.rgb = OWL_ORANGE_500
        para.add_run(". Estimated per-case savings: ").font.size = Pt(11)
        case_savings_run = para.add_run(fmt_currency(case_savings))
        case_savings_run.font.size = Pt(11)
        case_savings_run.bold = True
        case_savings_run.font.color.rgb = OWL_PURPLE_500
        para.add_run(f" compared to GPU server running for {case_duration_text}.").font.size = Pt(11)
    else:
        para = doc.add_paragraph()
        para.add_run(f"Per-case cost ({case_duration_text}, {case_documents:,} documents): ").font.size = Pt(11)
        case_cost_run = para.add_run(fmt_currency(preferred_case_cost))
        case_cost_run.font.size = Pt(11)
        case_cost_run.bold = True
        case_cost_run.font.color.rgb = OWL_ORANGE_500
        para.add_run(".").font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Pros and Cons
    sol_comp_heading = doc.add_heading('Solution Comparison', 2)
    sol_comp_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    # GPU Server Pros/Cons
    doc.add_heading('GPU Server (Ollama)', 3)
    
    doc.add_paragraph('Advantages:', style='List Bullet')
    for item in [
        "No API usage fees",
        "Complete data privacy (local processing)",
        "No external service dependencies",
        "Predictable monthly costs",
        "No rate limits"
    ]:
        para = doc.add_paragraph(item, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph('Considerations:', style='List Bullet')
    for item in [
        "Higher fixed costs ($1,440/month)",
        "Slower processing times",
        "Lower extraction quality",
        "Requires infrastructure management",
        "24/7 server operation required"
    ]:
        para = doc.add_paragraph(item, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()  # Spacing
    
    # Non-GPU Server Pros/Cons
    doc.add_heading('Non-GPU Server (OpenAI)', 3)
    
    doc.add_paragraph('Advantages:', style='List Bullet')
    for item in [
        f"Lower total costs ({fmt_currency(data.get('comparison', {}).get('openai_monthly', 301))}/month)",
        "Faster processing times",
        "Higher quality extraction",
        "Access to latest AI models",
        "No infrastructure management",
        "Better scalability"
    ]:
        para = doc.add_paragraph(item, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph('Considerations:', style='List Bullet')
    for item in [
        "API costs scale with volume",
        "Data processed externally",
        "Subject to API rate limits",
        "Variable costs based on usage"
    ]:
        para = doc.add_paragraph(item, style='List Bullet 2')
        para.paragraph_format.left_indent = Inches(0.5)
    
    add_page_break(doc)
    
    # SECTION 2: TECHNICAL BREAKDOWN
    heading2 = doc.add_heading('SECTION 2: TECHNICAL PERFORMANCE BREAKDOWN', 1)
    heading2.runs[0].font.color.rgb = OWL_BLUE_900
    
    # Overview section
    overview_heading = doc.add_heading('Model Overview & Selection Guidance', 2)
    overview_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    overview_para = doc.add_paragraph()
    overview_para.add_run(
        f"Six OpenAI models were evaluated for entity and relationship extraction from documents. "
        f"Each model represents a different balance between cost, speed, and extraction quality. "
        f"Metrics below are provided at three levels: per chunk (basic processing unit), per page "
        f"(assuming {data['avg_pages_per_document']} pages per document), and per document "
        f"(assuming {data['avg_chunks_per_document']:.1f} chunks per document)."
    )
    overview_para.runs[0].font.size = Pt(11)
    overview_para.paragraph_format.line_spacing = 1.5
    overview_para.paragraph_format.space_after = Pt(12)
    
    # Model selection guidance
    guidance_heading = doc.add_heading('Which Model to Use When', 3)
    guidance_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    # Find models by name for accurate guidance
    model_dict = {m['name'].lower(): m for m in data['openai_models']}
    
    def get_model_cost(name):
        return model_dict.get(name.lower(), {}).get('cost_per_chunk', 0)
    
    guidance_items = []
    
    # Investigation Type-Based Recommendations
    if 'gpt-4o-mini' in model_dict:
        guidance_items.append((
            "High-Volume, Cost-Sensitive Cases",
            f"Use gpt-4o-mini (${get_model_cost('gpt-4o-mini'):.6f}/chunk) for maximum cost efficiency. "
            f"Ideal for: routine due diligence, standard background checks, high-volume document processing where basic entity extraction is sufficient. "
            f"Best for cases with 100+ documents/month where cost optimization is critical."
        ))
    
    if 'gpt-4o' in model_dict or 'gpt-5-mini' in model_dict:
        gpt4o_cost = get_model_cost('gpt-4o')
        gpt5mini_cost = get_model_cost('gpt-5-mini')
        cost_text = []
        if gpt4o_cost > 0:
            cost_text.append(f"gpt-4o (${gpt4o_cost:.6f}/chunk)")
        if gpt5mini_cost > 0:
            cost_text.append(f"gpt-5-mini (${gpt5mini_cost:.6f}/chunk)")
        guidance_items.append((
            "Standard Fraud Investigations & Corporate Due Diligence",
            f"Use {' or '.join(cost_text)} for balanced quality-to-cost ratio. "
            f"gpt-5-mini offers better extraction with minimal cost increase. "
            f"Ideal for: financial fraud cases, corporate investigations, asset tracing, standard compliance reviews. "
            f"These models provide reliable entity and relationship extraction without premium pricing."
        ))
    
    if 'gpt-4.1' in model_dict:
        guidance_items.append((
            "Complex Fraud Cases & Enhanced Due Diligence",
            f"Use gpt-4.1 (${get_model_cost('gpt-4.1'):.6f}/chunk) when you need better extraction than GPT-4o "
            f"but want to avoid GPT-5 series pricing. "
            f"Ideal for: sophisticated fraud schemes, multi-party investigations, cases requiring detailed entity recognition. "
            f"Offers improved accuracy over GPT-4o while maintaining competitive speed and cost."
        ))
    
    if 'gpt-5.1' in model_dict:
        guidance_items.append((
            "RICO Cases & Complex Organizational Investigations",
            f"Use gpt-5.1 (${get_model_cost('gpt-5.1'):.6f}/chunk) for investigations requiring detailed relationship mapping and organizational structure analysis. "
            f"Provides excellent entity-relationship extraction without the maximum cost of gpt-5.2. "
            f"Ideal for: RICO investigations, organized crime cases, complex corporate structures, multi-entity fraud networks, "
            f"cases where understanding relationships between entities is critical. The model's superior relationship extraction "
            f"helps map intricate organizational hierarchies and connection patterns."
        ))
    
    if 'gpt-5.2' in model_dict:
        guidance_items.append((
            "Terrorism Investigations & Critical National Security Cases",
            f"Use gpt-5.2 (${get_model_cost('gpt-5.2'):.6f}/chunk) for high-stakes investigations where missing entities could have serious consequences. "
            f"Highest quality but slowest and most expensive. "
            f"Ideal for: terrorism investigations, national security cases, critical fraud investigations with high financial impact, "
            f"cases involving intricate corporate structures with shell companies, investigations where comprehensive analysis is paramount. "
            f"The model's maximum extraction depth (42 entities, 56 relationships per chunk) ensures no critical information is overlooked. "
            f"Reserve for cases where the cost of missing entities far exceeds the processing cost."
        ))
    
    # Add general guidance
    guidance_items.append((
        "General Selection Guidelines",
        f"Consider your investigation type: Standard fraud and due diligence cases typically benefit from gpt-4o or gpt-5-mini. "
        f"RICO and complex organizational cases require gpt-5.1 for relationship mapping. Terrorism and critical national security "
        f"cases demand gpt-5.2's maximum extraction capability. Always balance extraction quality needs against processing volume and budget constraints."
    ))
    
    for title, description in guidance_items:
        # Title
        item_title = doc.add_paragraph(title)
        item_title.runs[0].font.size = Pt(11)
        item_title.runs[0].font.color.rgb = OWL_BLUE_900
        item_title.runs[0].bold = True
        item_title.paragraph_format.space_before = Pt(6)
        
        # Description
        item_desc = doc.add_paragraph(description)
        item_desc.runs[0].font.size = Pt(10)
        item_desc.paragraph_format.left_indent = Inches(0.25)
        item_desc.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph()  # Extra spacing
    add_page_break(doc)
    
    # Detailed metrics intro
    metrics_intro = doc.add_heading('Detailed Model Performance', 2)
    metrics_intro.runs[0].font.color.rgb = OWL_BLUE_800
    
    intro_para = doc.add_paragraph()
    intro_para.add_run(
        "The following sections provide detailed performance metrics, model descriptions, and behavioral explanations "
        "for each tested model. Use this information to make informed decisions based on your specific case requirements."
    )
    intro_para.runs[0].font.size = Pt(11)
    intro_para.paragraph_format.space_after = Pt(12)
    
    # Technical breakdown for each model
    for model in data['openai_models']:
        model_heading = doc.add_heading(model['name'], 2)
        model_heading.runs[0].font.color.rgb = OWL_BLUE_800
        
        # Model description and behavior explanation
        description, behavior = get_model_description(model['name'], model)
        
        desc_para = doc.add_paragraph()
        desc_para.add_run("Description: ").font.size = Pt(11)
        desc_para.add_run(description).font.size = Pt(11)
        desc_para.runs[1].italic = True
        desc_para.paragraph_format.space_after = Pt(8)
        
        behavior_para = doc.add_paragraph()
        behavior_para.add_run("Why This Model Behaves This Way: ").font.size = Pt(11)
        behavior_para.runs[0].font.color.rgb = OWL_BLUE_800
        behavior_para.runs[0].bold = True
        behavior_para.add_run(behavior).font.size = Pt(10)
        behavior_para.paragraph_format.left_indent = Inches(0.15)
        behavior_para.paragraph_format.line_spacing = 1.4
        behavior_para.paragraph_format.space_after = Pt(12)
        
        metrics = calculate_metrics(data, model)
        
        # Create technical table
        tech_table = doc.add_table(rows=5, cols=4)
        tech_table.style = 'Light Grid Accent 1'
        tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header = tech_table.rows[0].cells
        header[0].text = 'Metric'
        header[1].text = 'Per Chunk'
        header[2].text = 'Per Page'
        header[3].text = 'Per Document'
        
        for cell in header:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1d4d76')  # owl-blue-800
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Cost row
        row1 = tech_table.rows[1].cells
        row1[0].text = 'Cost'
        row1[0].paragraphs[0].runs[0].font.bold = True
        row1[1].text = fmt_currency(metrics['per_chunk']['cost'])
        row1[2].text = fmt_currency(metrics['per_page']['cost'])
        row1[3].text = fmt_currency(metrics['per_document']['cost'])
        
        # Time row
        row2 = tech_table.rows[2].cells
        row2[0].text = 'Processing Time'
        row2[0].paragraphs[0].runs[0].font.bold = True
        row2[1].text = fmt_time(metrics['per_chunk']['time_seconds'])
        row2[2].text = fmt_time(metrics['per_page']['time_seconds'])
        row2[3].text = fmt_time(metrics['per_document']['time_seconds'])
        
        # Entities row
        row3 = tech_table.rows[3].cells
        row3[0].text = 'Entities Extracted'
        row3[0].paragraphs[0].runs[0].font.bold = True
        row3[1].text = str(int(metrics['per_chunk']['entities']))
        row3[2].text = str(int(metrics['per_page']['entities']))
        row3[3].text = str(int(metrics['per_document']['entities']))
        
        # Relationships row
        row4 = tech_table.rows[4].cells
        row4[0].text = 'Relationships Extracted'
        row4[0].paragraphs[0].runs[0].font.bold = True
        row4[1].text = str(int(metrics['per_chunk']['relationships']))
        row4[2].text = str(int(metrics['per_page']['relationships']))
        row4[3].text = str(int(metrics['per_document']['relationships']))
        
        # Center align all data cells
        for row in tech_table.rows[1:]:
            for cell in row.cells[1:]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Token usage
        para = doc.add_paragraph()
        para.add_run(f"Token Usage: ").font.bold = True
        para.add_run(f"{model.get('tokens', 0):,} tokens per chunk")
        para.runs[0].font.size = Pt(10)
        para.runs[1].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing between models
    
    # Quick Comparison Table
    quick_comp_heading = doc.add_heading('Quick Comparison - All Models', 2)
    quick_comp_heading.runs[0].font.color.rgb = OWL_BLUE_800
    
    comp_table = doc.add_table(rows=len(data['openai_models']) + 1, cols=5)
    comp_table.style = 'Light Grid Accent 1'
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    comp_header = comp_table.rows[0].cells
    comp_header[0].text = 'Model'
    comp_header[1].text = 'Cost/Doc'
    comp_header[2].text = 'Time/Doc'
    comp_header[3].text = 'Entities/Doc'
    comp_header[4].text = 'Relationships/Doc'
    
    for cell in comp_header:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1a365d')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for idx, model in enumerate(data['openai_models']):
        metrics = calculate_metrics(data, model)
        row = comp_table.rows[idx + 1].cells
        row[0].text = model['name']
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = fmt_currency(metrics['per_document']['cost'])
        row[2].text = fmt_time(metrics['per_document']['time_seconds'])
        row[3].text = str(int(metrics['per_document']['entities']))
        row[4].text = str(int(metrics['per_document']['relationships']))
        
        # Center align data
        for cell in row[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Footer
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Owl Consultancy Group - Infrastructure Cost Analysis Report')
    run.font.bold = True
    run.font.size = Pt(10)
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f"Generated on {data.get('generated_date', datetime.now().strftime('%B %d, %Y at %I:%M %p'))}")
    para.runs[0].font.size = Pt(9)
    para.runs[0].font.color.rgb = OWL_GRAY_500
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(
        f"This analysis is based on current pricing and performance metrics. Actual costs may vary based on "
        f"specific usage patterns, document complexity, and pricing changes. All metrics assume average "
        f"document size of {data['avg_pages_per_document']} pages and {data['avg_chunks_per_document']:.1f} chunks per document."
    )
    para.runs[0].font.size = Pt(8)
    para.runs[0].font.color.rgb = OWL_GRAY_500
    
    # Save document
    doc.save(output_path)
    return True


def main():
    """Main function to generate DOCX report."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx not available. Install with: pip install python-docx")
        return False
    
    txt_path = PROJECT_ROOT / "ingestion" / "data" / "cost_analysis_report.txt"
    docx_path = PROJECT_ROOT / "ingestion" / "data" / "cost_analysis_report.docx"
    
    if not txt_path.exists():
        print(f"Error: Text report not found at {txt_path}")
        return False
    
    print(f"Parsing text report: {txt_path}")
    data = parse_text_report(txt_path)
    
    print(f"Found {len(data['openai_models'])} OpenAI models")
    print(f"Generating DOCX report: {docx_path}")
    
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        generate_docx(data, docx_path)
        print(f"✓ DOCX report generated: {docx_path}")
        return True
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
