from __future__ import annotations

import csv
import html
import io
import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Literal

from postgres.models.agent import AgentArtifactRecord
from services.export_common import (
    AIDisclosureLevel,
    ExportMetadata,
    generate_export_id,
    render_metadata_block_docx,
    render_metadata_block_html,
    render_metadata_csv_columns,
    render_metadata_csv_values,
    safe_filename,
)


AgentExportFormat = Literal["csv", "pdf", "docx"]


@dataclass(frozen=True)
class AgentArtifactExport:
    content: bytes
    filename: str
    media_type: str
    export_id: str = ""


def render_artifact_export(
    artifact: AgentArtifactRecord,
    export_format: AgentExportFormat,
    *,
    generated_by: str = "Unknown user",
) -> AgentArtifactExport:
    payload = artifact.payload or {}
    export_metadata = _build_agent_export_metadata(
        artifact_type=artifact.type,
        title=artifact.title,
        payload=payload,
        case_id=_artifact_case_id(artifact),
        generated_by=generated_by,
        extra_metadata=artifact.extra_metadata,
    )
    if export_format == "pdf":
        if artifact.type != "report":
            raise ValueError("PDF export is only supported for report artifacts")
        return render_report_pdf(
            title=artifact.title,
            payload=payload,
            export_metadata=export_metadata,
            extra_metadata=artifact.extra_metadata,
        )
    if export_format == "docx":
        if artifact.type != "report":
            raise ValueError("Word export is only supported for report artifacts")
        return render_report_docx(
            title=artifact.title,
            payload=payload,
            export_metadata=export_metadata,
            extra_metadata=artifact.extra_metadata,
        )
    if export_format != "csv":
        raise ValueError(f"Unsupported artifact export format: {export_format}")
    return render_artifact_csv(
        artifact_type=artifact.type,
        title=artifact.title,
        payload=payload,
        export_metadata=export_metadata,
        extra_metadata=artifact.extra_metadata,
    )


def _build_agent_export_metadata(
    *,
    artifact_type: str,
    title: str,
    payload: dict[str, Any],
    case_id: str,
    generated_by: str,
    extra_metadata: dict[str, Any] | None = None,
) -> ExportMetadata:
    scope = str(
        payload.get("scope")
        or (extra_metadata or {}).get("query")
        or (extra_metadata or {}).get("notes")
        or ""
    ).strip()
    return ExportMetadata(
        export_id=generate_export_id(),
        case_id=str(case_id or "unknown"),
        generated_at=datetime.now(timezone.utc),
        generated_by=generated_by or "Unknown user",
        filters_description=scope or None,
        scope_description=f"Agent {artifact_type or 'artifact'} export: {title or 'Untitled artifact'}.",
        ai_disclosure=AIDisclosureLevel.AI_GENERATED,
        source_citations=_source_citations_from_payload(payload, extra_metadata),
    )


def _artifact_case_id(artifact: AgentArtifactRecord) -> str:
    run = getattr(artifact, "run", None)
    if run is not None and getattr(run, "case_id", None):
        return str(run.case_id)
    thread = getattr(artifact, "thread", None)
    if thread is not None and getattr(thread, "case_id", None):
        return str(thread.case_id)
    return "unknown"


def _source_citations_from_payload(
    payload: dict[str, Any],
    extra_metadata: dict[str, Any] | None = None,
) -> tuple[str, ...]:
    citations: list[str] = []

    def add(text: str | None) -> None:
        clean = str(text or "").strip()
        if clean and clean not in citations:
            citations.append(clean)

    for key in ("source_citations", "citations", "sources", "source_references"):
        for item in _citation_items(payload.get(key)):
            add(item)

    for result_id in _as_strings(payload.get("source_result_ids")):
        add(f"Agent source result: {result_id}")
    for result_id in _as_strings((extra_metadata or {}).get("source_result_ids")):
        add(f"Agent source result: {result_id}")

    for embed in _dict_rows(payload.get("embedded_artifacts")):
        embed_title = str(embed.get("title") or embed.get("artifact_id") or "Embedded artifact")
        embed_metadata = embed.get("metadata") if isinstance(embed.get("metadata"), dict) else {}
        for result_id in _as_strings(embed_metadata.get("source_result_ids")):
            add(f"{embed_title}: Agent source result {result_id}")

    for section in _dict_rows(payload.get("sections")):
        for embed in _dict_rows(section.get("embeds")):
            if not embed.get("available", True):
                add(f"Unavailable embedded artifact: {embed.get('artifact_id') or embed.get('title') or 'unknown'}")
            embed_metadata = embed.get("metadata") if isinstance(embed.get("metadata"), dict) else {}
            for result_id in _as_strings(embed_metadata.get("source_result_ids")):
                title = str(embed.get("title") or embed.get("artifact_id") or "Embedded artifact")
                add(f"{title}: Agent source result {result_id}")

    return tuple(citations)


def _citation_items(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, dict):
        parts = [
            str(value.get(key) or "").strip()
            for key in (
                "title",
                "source",
                "source_file",
                "filename",
                "file",
                "document",
                "id",
                "result_id",
            )
            if str(value.get(key) or "").strip()
        ]
        page = str(value.get("page") or value.get("source_page") or "").strip()
        if page:
            parts.append(f"p.{page}")
        quote = str(value.get("quote") or value.get("excerpt") or "").strip()
        if quote:
            parts.append(f"excerpt: {quote[:160]}")
        return [" | ".join(parts)] if parts else []
    if isinstance(value, (list, tuple, set)):
        items: list[str] = []
        for item in value:
            items.extend(_citation_items(item))
        return items
    return [str(value)]


def _as_strings(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, (list, tuple, set)):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    return [text] if text else []


def render_artifact_csv(
    *,
    artifact_type: str,
    title: str,
    payload: dict[str, Any],
    case_id: str = "unknown",
    generated_by: str = "Unknown user",
    export_metadata: ExportMetadata | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> AgentArtifactExport:
    export_metadata = export_metadata or _build_agent_export_metadata(
        artifact_type=artifact_type,
        title=title,
        payload=payload,
        case_id=case_id,
        generated_by=generated_by,
        extra_metadata=extra_metadata,
    )
    rows = _rows_for_artifact(artifact_type, payload)
    columns = _columns_for_artifact(artifact_type, payload, rows)
    csv_text = _write_csv(columns, rows, export_metadata)
    filename = f"{safe_filename(title or 'agent-artifact', fallback='agent-artifact')}-{artifact_type}.csv"
    return AgentArtifactExport(
        content=csv_text.encode("utf-8-sig"),
        filename=filename,
        media_type="text/csv; charset=utf-8",
        export_id=export_metadata.export_id,
    )


def _rows_for_artifact(artifact_type: str, payload: dict[str, Any]) -> list[dict[str, Any]]:
    if artifact_type == "table":
        return _dict_rows(payload.get("rows"))
    if artifact_type == "chart":
        return _dict_rows(payload.get("rows"))
    if artifact_type == "map":
        return _dict_rows(payload.get("locations"))
    if artifact_type == "graph":
        return _graph_rows(payload)
    return _dict_rows(payload.get("rows") or payload.get("items") or payload.get("data"))


def _columns_for_artifact(
    artifact_type: str,
    payload: dict[str, Any],
    rows: list[dict[str, Any]],
) -> list[str]:
    if artifact_type == "table":
        explicit = [
            str(column.get("key"))
            for column in _dict_rows(payload.get("columns"))
            if column.get("key") is not None
        ]
        return _dedupe([*explicit, *_flattened_keys(rows)])
    if artifact_type == "chart":
        explicit = [
            str(column.get("key"))
            for column in _dict_rows(payload.get("columns"))
            if column.get("key") is not None
        ]
        chart_keys = [
            str(payload.get("x_key") or ""),
            str(payload.get("category_key") or ""),
            str(payload.get("value_key") or ""),
            *[str(key) for key in payload.get("y_keys") or []],
        ]
        return _dedupe([*chart_keys, *explicit, *_flattened_keys(rows)])

    preferred: dict[str, list[str]] = {
        "map": [
            "name",
            "key",
            "type",
            "latitude",
            "longitude",
            "address",
            "summary",
        ],
        "graph": [
            "row_type",
            "key",
            "name",
            "type",
            "source",
            "target",
            "summary",
            "date",
            "amount",
            "properties.date",
            "properties.amount",
            "properties.currency",
            "properties.detail",
            "properties.source_files",
        ],
    }
    if artifact_type in preferred:
        return _dedupe([*preferred[artifact_type], *[key for key in _flattened_keys(rows) if key in preferred[artifact_type]]])
    return _dedupe(_flattened_keys(rows))


def _graph_rows(payload: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for node in _dict_rows(payload.get("nodes")):
        rows.append({"row_type": "node", **node})
    for link in _dict_rows(payload.get("links")):
        rows.append({"row_type": "relationship", **link})
    return rows


def _write_csv(
    columns: list[str],
    rows: list[dict[str, Any]],
    export_metadata: ExportMetadata,
) -> str:
    if not columns:
        columns = ["value"]

    metadata_columns = render_metadata_csv_columns()
    metadata_values = render_metadata_csv_values(export_metadata)
    buffer = io.StringIO(newline="")
    writer = csv.DictWriter(
        buffer,
        fieldnames=[*columns, *metadata_columns],
        extrasaction="ignore",
        lineterminator="\n",
    )
    writer.writeheader()
    for row in rows:
        flattened = _flatten_dict(row)
        writer.writerow(
            {
                **{column: _cell_value(flattened.get(column)) for column in columns},
                **metadata_values,
            }
        )
    if not rows:
        writer.writerow(metadata_values)
    return buffer.getvalue()


def _flattened_keys(rows: list[dict[str, Any]]) -> list[str]:
    keys: list[str] = []
    for row in rows:
        for key in _flatten_dict(row):
            if key not in keys:
                keys.append(key)
    return keys


def _flatten_dict(value: dict[str, Any], prefix: str = "") -> dict[str, Any]:
    flattened: dict[str, Any] = {}
    for key, item in value.items():
        next_key = f"{prefix}.{key}" if prefix else str(key)
        if isinstance(item, dict):
            flattened.update(_flatten_dict(item, next_key))
        else:
            flattened[next_key] = item
    return flattened


def _cell_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def _dict_rows(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]


def _dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        if not value or value in seen:
            continue
        seen.add(value)
        result.append(value)
    return result


def render_report_pdf(
    *,
    title: str,
    payload: dict[str, Any],
    export_metadata: ExportMetadata | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> AgentArtifactExport:
    try:
        from weasyprint import HTML
    except Exception as exc:
        raise ValueError(f"PDF export is unavailable: {exc}") from exc

    export_metadata = export_metadata or _build_agent_export_metadata(
        artifact_type="report",
        title=title,
        payload=payload,
        case_id="unknown",
        generated_by="Unknown user",
        extra_metadata=extra_metadata,
    )
    html_text = _report_html(
        title=title,
        payload=payload,
        export_metadata=export_metadata,
        extra_metadata=extra_metadata,
    )
    return AgentArtifactExport(
        content=HTML(string=html_text).write_pdf(),
        filename=f"{safe_filename(title or 'agent-report', fallback='agent-report')}-report.pdf",
        media_type="application/pdf",
        export_id=export_metadata.export_id,
    )


def render_report_docx(
    *,
    title: str,
    payload: dict[str, Any],
    export_metadata: ExportMetadata | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> AgentArtifactExport:
    try:
        from docx import Document
    except Exception as exc:
        raise ValueError(f"Word export is unavailable: {exc}") from exc

    export_metadata = export_metadata or _build_agent_export_metadata(
        artifact_type="report",
        title=title,
        payload=payload,
        case_id="unknown",
        generated_by="Unknown user",
        extra_metadata=extra_metadata,
    )
    document = Document()
    document.add_heading(str(payload.get("title") or title or "Agent report"), level=0)
    render_metadata_block_docx(document, export_metadata)
    purpose = str(payload.get("purpose") or "").strip()
    if purpose:
        document.add_paragraph(purpose)
    scope = str(payload.get("scope") or "").strip()
    if scope:
        document.add_heading("Scope", level=1)
        document.add_paragraph(scope)
    included = [str(item) for item in payload.get("included_items") or [] if str(item).strip()]
    if included:
        document.add_heading("Included", level=1)
        for item in included:
            document.add_paragraph(item, style="List Bullet")

    for section in _dict_rows(payload.get("sections")):
        heading = str(section.get("heading") or "Section")
        level = max(1, min(int(section.get("level") or 2), 3))
        document.add_heading(heading, level=level)
        _add_markdownish_docx(document, str(section.get("content") or ""))
        for embed in _dict_rows(section.get("embeds")):
            _add_embed_docx(document, embed)

    open_questions = [str(item) for item in payload.get("open_questions") or [] if str(item).strip()]
    if open_questions:
        document.add_heading("Open Questions", level=1)
        for item in open_questions:
            document.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return AgentArtifactExport(
        content=buffer.getvalue(),
        filename=f"{safe_filename(title or 'agent-report', fallback='agent-report')}-report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        export_id=export_metadata.export_id,
    )


def _report_html(
    *,
    title: str,
    payload: dict[str, Any],
    export_metadata: ExportMetadata | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> str:
    export_metadata = export_metadata or _build_agent_export_metadata(
        artifact_type="report",
        title=title,
        payload=payload,
        case_id="unknown",
        generated_by="Unknown user",
        extra_metadata=extra_metadata,
    )
    report_title = str(payload.get("title") or title or "Agent report")
    sections = _dict_rows(payload.get("sections"))
    included = [str(item) for item in payload.get("included_items") or [] if str(item).strip()]
    open_questions = [str(item) for item in payload.get("open_questions") or [] if str(item).strip()]
    purpose = str(payload.get("purpose") or "").strip()
    scope = str(payload.get("scope") or "").strip()
    audience = str(payload.get("audience") or "").strip()

    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        "<style>",
        "body{font-family:Arial,sans-serif;color:#111827;margin:36px;line-height:1.45}",
        "h1{font-size:28px;margin:0 0 8px}h2{font-size:18px;margin:24px 0 8px}h3{font-size:15px;margin:18px 0 6px}",
        ".meta{color:#4b5563;font-size:12px;margin-bottom:18px}.scope{border-left:3px solid #d97706;padding-left:12px;color:#374151}",
        "ul{margin-top:6px}.section{break-inside:avoid}.embed{border:1px solid #d1d5db;border-radius:6px;padding:10px;margin:10px 0 14px}",
        ".embed-title{font-size:12px;text-transform:uppercase;letter-spacing:.04em;color:#6b7280;margin-bottom:6px}",
        "table{border-collapse:collapse;width:100%;font-size:11px;margin-top:6px}th,td{border:1px solid #d1d5db;padding:5px;text-align:left;vertical-align:top}th{background:#f3f4f6}",
        ".graph-list{font-size:12px;color:#374151}.muted{color:#6b7280}",
        "</style></head><body>",
        f"<h1>{html.escape(report_title)}</h1>",
        render_metadata_block_html(export_metadata),
    ]
    if audience:
        parts.append(f"<div class='meta'>Audience: {html.escape(audience)}</div>")
    if purpose:
        parts.append(f"<p>{html.escape(purpose)}</p>")
    if scope:
        parts.append(f"<h2>Scope</h2><p class='scope'>{html.escape(scope)}</p>")
    if included:
        parts.append("<h2>Included</h2><ul>")
        parts.extend(f"<li>{html.escape(item)}</li>" for item in included)
        parts.append("</ul>")

    for section in sections:
        level = max(2, min(int(section.get("level") or 2), 3))
        heading_tag = f"h{level}"
        parts.append("<div class='section'>")
        parts.append(f"<{heading_tag}>{html.escape(str(section.get('heading') or 'Section'))}</{heading_tag}>")
        parts.append(_markdownish_to_html(str(section.get("content") or "")))
        for embed in _dict_rows(section.get("embeds")):
            parts.append(_embed_html(embed))
        parts.append("</div>")

    if open_questions:
        parts.append("<h2>Open Questions</h2><ul>")
        parts.extend(f"<li>{html.escape(item)}</li>" for item in open_questions)
        parts.append("</ul>")

    citations = list(export_metadata.source_citations)
    if citations:
        parts.append("<h2>Source Citations</h2><ol>")
        parts.extend(f"<li>{html.escape(citation)}</li>" for citation in citations[:50])
        if len(citations) > 50:
            parts.append(f"<li>{len(citations) - 50} additional source citation(s) omitted.</li>")
        parts.append("</ol>")

    parts.append("</body></html>")
    return "".join(parts)


def _markdownish_to_html(content: str) -> str:
    lines = [line.rstrip() for line in content.splitlines()]
    parts: list[str] = []
    list_open = False
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            parts.append(f"<p>{html.escape(' '.join(paragraph))}</p>")
            paragraph = []

    def close_list() -> None:
        nonlocal list_open
        if list_open:
            parts.append("</ul>")
            list_open = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            close_list()
            continue
        if stripped.startswith(("- ", "* ")):
            flush_paragraph()
            if not list_open:
                parts.append("<ul>")
                list_open = True
            parts.append(f"<li>{html.escape(stripped[2:].strip())}</li>")
            continue
        close_list()
        paragraph.append(stripped)

    flush_paragraph()
    close_list()
    return "".join(parts)


def _embed_html(embed: dict[str, Any]) -> str:
    title = html.escape(str(embed.get("title") or "Embedded artifact"))
    caption = html.escape(str(embed.get("caption") or ""))
    artifact_type = str(embed.get("type") or "unknown")
    data = embed.get("data") if isinstance(embed.get("data"), dict) else {}
    if not embed.get("available", True):
        return f"<div class='embed'><div class='embed-title'>{title}</div><p class='muted'>Referenced artifact is not available in this report export.</p></div>"
    if artifact_type == "table":
        rows = _dict_rows(data.get("rows"))[:20]
        columns = _columns_for_artifact("table", data, rows)[:8]
        return (
            f"<div class='embed'><div class='embed-title'>Table: {title}</div>"
            + (f"<p class='muted'>{caption}</p>" if caption else "")
            + _html_table(columns, rows)
            + "</div>"
        )
    if artifact_type == "graph":
        nodes = _dict_rows(data.get("nodes"))[:20]
        links = _dict_rows(data.get("links"))[:20]
        node_names = ", ".join(str(node.get("name") or node.get("key") or "") for node in nodes[:10])
        return (
            f"<div class='embed'><div class='embed-title'>Graph: {title}</div>"
            + (f"<p class='muted'>{caption}</p>" if caption else "")
            + f"<p class='graph-list'>{len(nodes)} shown node(s), {len(links)} shown relationship(s).</p>"
            + (f"<p class='graph-list'>Key nodes: {html.escape(node_names)}</p>" if node_names else "")
            + "</div>"
        )
    if artifact_type == "chart":
        rows = _dict_rows(data.get("rows"))[:20]
        columns = _columns_for_artifact("chart", data, rows)[:8]
        chart_type = html.escape(str(data.get("chart_type") or "chart").replace("_", " "))
        series = ", ".join(str(item.get("label") or item.get("key") or "") for item in _dict_rows(data.get("series"))[:6])
        return (
            f"<div class='embed'><div class='embed-title'>Chart: {title}</div>"
            + (f"<p class='muted'>{caption}</p>" if caption else "")
            + f"<p class='graph-list'>Type: {chart_type}. Rows: {len(rows)}."
            + (f" Series: {html.escape(series)}." if series else "")
            + "</p>"
            + _html_table(columns, rows)
            + "</div>"
        )
    return f"<div class='embed'><div class='embed-title'>{title}</div><p class='muted'>Unsupported embedded artifact type: {html.escape(artifact_type)}</p></div>"


def _html_table(columns: list[str], rows: list[dict[str, Any]]) -> str:
    if not rows:
        return "<p class='muted'>No rows.</p>"
    header = "".join(f"<th>{html.escape(column)}</th>" for column in columns)
    body = []
    for row in rows:
        flattened = _flatten_dict(row)
        body.append("<tr>" + "".join(f"<td>{html.escape(_cell_value(flattened.get(column)))}</td>" for column in columns) + "</tr>")
    return f"<table><thead><tr>{header}</tr></thead><tbody>{''.join(body)}</tbody></table>"


def _add_markdownish_docx(document: Any, content: str) -> None:
    for block in re.split(r"\n\s*\n", content.strip()):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        if all(line.startswith(("- ", "* ")) for line in lines):
            for line in lines:
                document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(" ".join(lines))


def _add_embed_docx(document: Any, embed: dict[str, Any]) -> None:
    title = str(embed.get("title") or "Embedded artifact")
    artifact_type = str(embed.get("type") or "unknown")
    document.add_paragraph(f"Embedded {artifact_type}: {title}")
    caption = str(embed.get("caption") or "").strip()
    if caption:
        document.add_paragraph(caption)
    data = embed.get("data") if isinstance(embed.get("data"), dict) else {}
    if artifact_type == "table":
        rows = _dict_rows(data.get("rows"))[:20]
        columns = _columns_for_artifact("table", data, rows)[:8]
        if rows and columns:
            table = document.add_table(rows=1, cols=len(columns))
            header_cells = table.rows[0].cells
            for index, column in enumerate(columns):
                header_cells[index].text = column
            for row in rows:
                cells = table.add_row().cells
                flattened = _flatten_dict(row)
                for index, column in enumerate(columns):
                    cells[index].text = _cell_value(flattened.get(column))
    elif artifact_type == "graph":
        nodes = _dict_rows(data.get("nodes"))
        links = _dict_rows(data.get("links"))
        document.add_paragraph(f"{len(nodes)} node(s), {len(links)} relationship(s)")
    elif artifact_type == "chart":
        chart_type = str(data.get("chart_type") or "chart").replace("_", " ")
        rows = _dict_rows(data.get("rows"))[:20]
        columns = _columns_for_artifact("chart", data, rows)[:8]
        document.add_paragraph(f"Chart type: {chart_type}; {len(rows)} source row(s)")
        if rows and columns:
            table = document.add_table(rows=1, cols=len(columns))
            header_cells = table.rows[0].cells
            for index, column in enumerate(columns):
                header_cells[index].text = column
            for row in rows:
                cells = table.add_row().cells
                flattened = _flatten_dict(row)
                for index, column in enumerate(columns):
                    cells[index].text = _cell_value(flattened.get(column))
