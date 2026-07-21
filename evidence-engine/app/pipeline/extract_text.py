import base64
import csv
import io
import json
import logging
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import docx
import openpyxl

from app.config import settings
from app.pipeline.pdf_extraction import PdfProgressCallback, extract_pdf
from app.services.openai_client import chat_completion, transcribe_audio
from app.utils.text_sanitize import sanitize_json, sanitize_text

logger = logging.getLogger(__name__)

AUDIO_EXTENSIONS = {
    ".mp3", ".m4a", ".wav", ".mpeg", ".mpga", ".ogg", ".flac",
}
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".webm"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif"}
MAX_WHISPER_SIZE = 25 * 1024 * 1024  # 25 MB
MIN_AUDIO_TRANSCRIPTION_SEGMENT_SECONDS = 60
TRANSCRIPTION_PROMPT_CONTEXT_CHARS = 1200


@dataclass
class ExtractedDocument:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    tables: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(file_path: str) -> ExtractedDocument:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables: list[str] = []
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))

    return ExtractedDocument(
        text="\n\n".join(paragraphs),
        metadata={
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        },
        tables=tables,
    )


# ---------------------------------------------------------------------------
# XLSX / XLS
# ---------------------------------------------------------------------------

def _extract_xlsx(file_path: str) -> ExtractedDocument:
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    tables: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            tables.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))

    sheet_count = len(wb.sheetnames)
    wb.close()

    return ExtractedDocument(
        text="",
        metadata={"file_type": "xlsx", "sheet_count": sheet_count},
        tables=tables,
    )


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def _extract_csv(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")

    reader = csv.DictReader(io.StringIO(raw))
    rows_data: list[dict[str, str]] = []
    for row in reader:
        rows_data.append(row)

    text = json.dumps(rows_data, indent=2, default=str) if rows_data else ""

    return ExtractedDocument(
        text="",
        metadata={
            "file_type": "csv",
            "row_count": len(rows_data),
            "columns": list(rows_data[0].keys()) if rows_data else [],
        },
        tables=[text] if text else [],
    )


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def _extract_html(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")

    try:
        from lxml import html as lxml_html

        doc = lxml_html.fromstring(raw)
        # Remove script and style elements
        for el in doc.iter("script", "style"):
            el.drop_tree()
        text = doc.text_content()
    except Exception:
        # Fallback: strip tags with regex
        text = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", raw, flags=re.IGNORECASE)
        text = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"&[a-zA-Z]+;", " ", text)

    # Normalise whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return ExtractedDocument(
        text=text,
        metadata={"file_type": "html"},
    )


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def _extract_markdown(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")

    try:
        from markdown_it import MarkdownIt

        md = MarkdownIt()
        html = md.render(raw)
        # Strip HTML tags from rendered output
        text = re.sub(r"<[^>]+>", " ", html)
    except Exception:
        # Fallback: strip markdown syntax with regex
        text = raw
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)  # headings
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
        text = re.sub(r"`{1,3}[^`]*`{1,3}", "", text)  # inline/block code
        text = re.sub(r"!?\[([^\]]*)\]\([^)]+\)", r"\1", text)  # links/images
        text = re.sub(r"^[-*+]\s+", "", text, flags=re.MULTILINE)  # list markers
        text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)  # blockquotes

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return ExtractedDocument(
        text=text,
        metadata={"file_type": "markdown"},
    )


# ---------------------------------------------------------------------------
# Audio
# ---------------------------------------------------------------------------

class AudioTranscriptionError(RuntimeError):
    """Raised when an audio segment cannot be transcribed."""


def _configured_audio_segment_seconds() -> int:
    return max(
        MIN_AUDIO_TRANSCRIPTION_SEGMENT_SECONDS,
        int(settings.audio_transcription_segment_seconds or 240),
    )


def _configured_audio_max_single_seconds() -> int:
    return max(
        MIN_AUDIO_TRANSCRIPTION_SEGMENT_SECONDS,
        int(settings.audio_transcription_max_single_seconds or 240),
    )


def _probe_media_duration_seconds(file_path: str) -> float | None:
    """Return media duration using ffprobe, or None when probing fails."""
    try:
        result = subprocess.run(
            [
                "ffprobe",
                "-v",
                "quiet",
                "-print_format",
                "json",
                "-show_format",
                file_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0 or not result.stdout:
            return None
        info = json.loads(result.stdout)
        duration = (info.get("format") or {}).get("duration")
        if duration is None:
            return None
        parsed = float(duration)
        return parsed if parsed > 0 else None
    except Exception as exc:
        logger.warning("Audio duration probe failed for %s: %s", file_path, exc)
        return None


def _split_audio_segments(
    file_path: str,
    output_dir: str,
    segment_seconds: int,
) -> list[Path]:
    """Split audio into normalized MP3 chunks for transcription."""
    segment_pattern = os.path.join(output_dir, "segment_%03d.mp3")
    try:
        subprocess.run(
            [
                "ffmpeg",
                "-hide_banner",
                "-loglevel",
                "error",
                "-y",
                "-i",
                file_path,
                "-map",
                "0:a:0",
                "-vn",
                "-ac",
                "1",
                "-ar",
                "16000",
                "-c:a",
                "libmp3lame",
                "-b:a",
                "64k",
                "-f",
                "segment",
                "-segment_time",
                str(segment_seconds),
                "-reset_timestamps",
                "1",
                segment_pattern,
            ],
            check=True,
            capture_output=True,
            timeout=1800,
        )
    except subprocess.CalledProcessError as exc:
        stderr = (
            exc.stderr.decode("utf-8", errors="replace")
            if isinstance(exc.stderr, bytes)
            else (exc.stderr or "")
        )
        raise AudioTranscriptionError(
            f"Audio segmentation failed: {stderr.strip() or exc}"
        ) from exc

    segments = sorted(Path(output_dir).glob("segment_*.mp3"))
    if not segments:
        raise AudioTranscriptionError("Audio segmentation produced no chunks")
    return segments


def _is_input_too_large_error(exc: Exception) -> bool:
    candidates: list[str] = []
    for attr in ("code", "message", "param"):
        value = getattr(exc, attr, None)
        if value:
            candidates.append(str(value))
    body = getattr(exc, "body", None)
    if isinstance(body, dict):
        candidates.extend(str(value) for value in body.values() if value is not None)
    elif body:
        candidates.append(str(body))
    candidates.append(str(exc))
    text = " ".join(candidates).lower()
    return "input_too_large" in text or "too large for this model" in text


def _build_transcription_prompt(previous_transcript: str | None) -> str | None:
    if not previous_transcript:
        return None
    tail = previous_transcript.strip()[-TRANSCRIPTION_PROMPT_CONTEXT_CHARS:]
    if not tail:
        return None
    return (
        "The previous audio segment ended with this transcript. Continue "
        "transcribing the same conversation, preserving names, punctuation, "
        "and context.\n\nPrevious segment ending:\n"
        f"{tail}"
    )


async def _transcribe_audio_segments(
    segments: list[Path],
    segment_seconds: int,
    transcripts: list[str],
    stats: dict[str, int],
) -> None:
    total_segments = len(segments)
    for index, segment in enumerate(segments, start=1):
        prompt = _build_transcription_prompt(transcripts[-1] if transcripts else None)
        try:
            transcript = await transcribe_audio(str(segment), prompt=prompt)
            stats["segment_count"] = stats.get("segment_count", 0) + 1
        except Exception as exc:
            if (
                _is_input_too_large_error(exc)
                and segment_seconds > MIN_AUDIO_TRANSCRIPTION_SEGMENT_SECONDS
            ):
                retry_seconds = max(
                    MIN_AUDIO_TRANSCRIPTION_SEGMENT_SECONDS,
                    segment_seconds // 2,
                )
                logger.warning(
                    "Audio segment %s/%s was too large for transcription; retrying with %ss chunks",
                    index,
                    total_segments,
                    retry_seconds,
                )
                with tempfile.TemporaryDirectory() as retry_dir:
                    retry_segments = _split_audio_segments(
                        str(segment),
                        retry_dir,
                        retry_seconds,
                    )
                    try:
                        await _transcribe_audio_segments(
                            retry_segments,
                            retry_seconds,
                            transcripts,
                            stats,
                        )
                    except Exception as retry_exc:
                        raise AudioTranscriptionError(
                            "Audio transcription failed on segment "
                            f"{index}/{total_segments} after retrying with "
                            f"{retry_seconds}s chunks: {retry_exc}"
                        ) from retry_exc
                continue

            raise AudioTranscriptionError(
                f"Audio transcription failed on segment {index}/{total_segments}: {exc}"
            ) from exc

        transcript = transcript.strip()
        if transcript:
            transcripts.append(transcript)


async def _extract_audio(file_path: str) -> ExtractedDocument:
    file_size = os.path.getsize(file_path)
    duration_seconds = _probe_media_duration_seconds(file_path)
    segment_seconds = _configured_audio_segment_seconds()
    max_single_seconds = _configured_audio_max_single_seconds()
    metadata: dict[str, Any] = {
        "file_type": "audio",
        "file_size": file_size,
    }
    if duration_seconds is not None:
        metadata["duration_seconds"] = duration_seconds

    should_segment = file_size > MAX_WHISPER_SIZE or (
        duration_seconds is not None and duration_seconds > max_single_seconds
    )
    if not should_segment:
        try:
            transcript = await transcribe_audio(file_path)
            metadata["segment_count"] = 1
            metadata["transcription"] = transcript
            return ExtractedDocument(text=transcript, metadata=metadata)
        except Exception as exc:
            if not _is_input_too_large_error(exc):
                raise
            logger.warning(
                "Single audio transcription request was too large; falling back to %ss chunks",
                segment_seconds,
            )

    stats = {"segment_count": 0}
    transcripts: list[str] = []
    with tempfile.TemporaryDirectory() as segments_dir:
        segments = _split_audio_segments(file_path, segments_dir, segment_seconds)
        await _transcribe_audio_segments(
            segments,
            segment_seconds,
            transcripts,
            stats,
        )

    transcription = "\n\n".join(transcripts)
    metadata["segment_count"] = stats["segment_count"]
    metadata["segment_seconds"] = segment_seconds
    metadata["transcription"] = transcription
    return ExtractedDocument(
        text=transcription,
        metadata=metadata,
    )


# ---------------------------------------------------------------------------
# Image (Tesseract OCR / OpenAI Vision)
# ---------------------------------------------------------------------------

def _extract_image_metadata(file_path: str) -> dict[str, Any]:
    """Extract EXIF metadata from an image file."""
    metadata: dict[str, Any] = {"file_type": "image"}
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS

        img = Image.open(file_path)
        metadata["width"] = img.width
        metadata["height"] = img.height
        metadata["format"] = img.format

        exif_data = img.getexif()
        if exif_data:
            for tag_id, value in exif_data.items():
                tag_name = TAGS.get(tag_id, str(tag_id))
                if tag_name in ("DateTime", "DateTimeOriginal", "Make", "Model"):
                    metadata[tag_name] = str(value)
        img.close()
    except Exception:
        pass
    return metadata


def _extract_text_tesseract(file_path: str) -> str:
    """OCR an image using Tesseract."""
    import pytesseract
    from PIL import Image

    img = Image.open(file_path)
    text = pytesseract.image_to_string(img, lang=settings.tesseract_lang)
    img.close()
    return text.strip()


async def _describe_image_vision(file_path: str, file_name: str) -> str:
    """Describe an image using OpenAI Vision API."""
    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")

    ext = Path(file_path).suffix.lower()
    mime_map = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".tiff": "image/tiff", ".tif": "image/tiff",
        ".bmp": "image/bmp",
    }
    mime = mime_map.get(ext, "image/jpeg")

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        f"This image is from a file called '{file_name}' uploaded as evidence "
                        "in an investigation. Describe everything you see in detail: people, "
                        "objects, text, locations, documents, licence plates, timestamps, "
                        "and any other relevant details. If there is text in the image, "
                        "transcribe it exactly."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{encoded}"},
                },
            ],
        }
    ]

    return await chat_completion(
        messages=messages,
        model=settings.openai_vision_model,
    )


async def _extract_image(file_path: str, file_name: str) -> ExtractedDocument:
    metadata = _extract_image_metadata(file_path)
    provider = settings.image_provider.lower()
    text = ""

    if provider == "openai":
        try:
            text = await _describe_image_vision(file_path, file_name)
            metadata["image_provider"] = "openai_vision"
        except Exception as e:
            logger.warning("OpenAI Vision failed, falling back to Tesseract: %s", e)
            try:
                text = _extract_text_tesseract(file_path)
                metadata["image_provider"] = "tesseract_fallback"
            except Exception:
                logger.error("Both image extraction methods failed for %s", file_name)
    else:
        try:
            text = _extract_text_tesseract(file_path)
            metadata["image_provider"] = "tesseract"
        except Exception as e:
            logger.warning("Tesseract failed, falling back to OpenAI Vision: %s", e)
            try:
                text = await _describe_image_vision(file_path, file_name)
                metadata["image_provider"] = "openai_vision_fallback"
            except Exception:
                logger.error("Both image extraction methods failed for %s", file_name)

    return ExtractedDocument(text=text, metadata=metadata)


# ---------------------------------------------------------------------------
# Video (FFmpeg frame extraction + Vision + Whisper)
# ---------------------------------------------------------------------------

def _get_video_metadata(file_path: str) -> dict[str, Any]:
    """Extract video metadata using ffprobe."""
    metadata: dict[str, Any] = {"file_type": "video"}
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "quiet", "-print_format", "json",
                "-show_format", "-show_streams", file_path,
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            info = json.loads(result.stdout)
            fmt = info.get("format", {})
            metadata["duration_seconds"] = float(fmt.get("duration", 0))
            metadata["format_name"] = fmt.get("format_name", "")
            for stream in info.get("streams", []):
                if stream.get("codec_type") == "video":
                    metadata["width"] = stream.get("width")
                    metadata["height"] = stream.get("height")
                    metadata["video_codec"] = stream.get("codec_name")
                elif stream.get("codec_type") == "audio":
                    metadata["has_audio"] = True
                    metadata["audio_codec"] = stream.get("codec_name")
    except Exception as e:
        logger.warning("ffprobe failed: %s", e)
    return metadata


def _extract_key_frames(
    file_path: str, output_dir: str
) -> list[dict[str, Any]]:
    """Extract key frames from a video at regular intervals."""
    interval = settings.video_frame_interval
    max_frames = settings.video_max_frames

    subprocess.run(
        [
            "ffmpeg", "-i", file_path,
            "-vf", f"fps=1/{interval}",
            "-frames:v", str(max_frames),
            "-q:v", "2",
            os.path.join(output_dir, "frame_%04d.jpg"),
        ],
        check=True, capture_output=True, timeout=300,
    )

    frames: list[dict[str, Any]] = []
    for i, frame_file in enumerate(sorted(Path(output_dir).glob("frame_*.jpg"))):
        timestamp = i * interval
        frames.append({
            "path": str(frame_file),
            "timestamp_seconds": timestamp,
            "timestamp_str": f"{timestamp // 60:02d}:{timestamp % 60:02d}",
        })
    return frames


def _extract_audio_track(file_path: str, output_dir: str) -> str | None:
    """Extract audio track from video as WAV."""
    audio_path = os.path.join(output_dir, "audio.wav")
    try:
        subprocess.run(
            [
                "ffmpeg", "-i", file_path,
                "-vn", "-acodec", "pcm_s16le",
                "-ar", "16000", "-ac", "1",
                audio_path,
            ],
            check=True, capture_output=True, timeout=300,
        )
        if os.path.exists(audio_path) and os.path.getsize(audio_path) > 0:
            return audio_path
    except Exception as e:
        logger.warning("Audio extraction from video failed: %s", e)
    return None


async def _extract_video(file_path: str, file_name: str) -> ExtractedDocument:
    metadata = _get_video_metadata(file_path)
    parts: list[str] = []

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Extract and transcribe audio track
        audio_path = _extract_audio_track(file_path, tmp_dir)
        if audio_path:
            try:
                audio_doc = await _extract_audio(audio_path)
                if audio_doc.text.strip():
                    parts.append("[Audio Transcription]\n" + audio_doc.text)
                    metadata["has_transcription"] = True
                    metadata["transcription"] = audio_doc.text
            except Exception as e:
                logger.warning("Video audio transcription failed: %s", e)

        # Extract key frames
        frames_dir = os.path.join(tmp_dir, "frames")
        os.makedirs(frames_dir, exist_ok=True)
        try:
            frames = _extract_key_frames(file_path, frames_dir)
            metadata["frame_count"] = len(frames)
        except Exception as e:
            logger.warning("Frame extraction failed: %s", e)
            frames = []

        # Describe frames with Vision API
        if frames:
            frame_descriptions: list[str] = []
            for frame in frames:
                try:
                    description = await _describe_image_vision(
                        frame["path"], f"{file_name} @ {frame['timestamp_str']}"
                    )
                    frame_descriptions.append(
                        f"[{frame['timestamp_str']}] {description}"
                    )
                except Exception as e:
                    logger.warning(
                        "Frame description failed at %s: %s",
                        frame["timestamp_str"], e,
                    )
            if frame_descriptions:
                parts.append(
                    "[Video Frame Descriptions]\n" + "\n\n".join(frame_descriptions)
                )

    text = "\n\n".join(parts)
    return ExtractedDocument(text=text, metadata=metadata)


# ---------------------------------------------------------------------------
# Main dispatcher
# ---------------------------------------------------------------------------

def _sanitize_extracted(doc: ExtractedDocument) -> ExtractedDocument:
    return ExtractedDocument(
        text=sanitize_text(doc.text),
        metadata=sanitize_json(doc.metadata) if doc.metadata else {},
        tables=[sanitize_text(t) for t in (doc.tables or [])],
    )


def get_transcription(doc: ExtractedDocument) -> str | None:
    """Return the full transcript when extraction produced one."""
    transcription = doc.metadata.get("transcription")
    if isinstance(transcription, str) and transcription.strip():
        return sanitize_text(transcription.strip())
    return None


async def extract_text(
    file_path: str,
    file_name: str,
    progress_callback: PdfProgressCallback | None = None,
) -> ExtractedDocument:
    ext = Path(file_name).suffix.lower()

    if ext == ".pdf":
        pdf = await extract_pdf(file_path, progress_callback=progress_callback)
        doc = ExtractedDocument(
            text=pdf.text,
            metadata=pdf.metadata,
            tables=pdf.tables,
        )
    elif ext in (".docx", ".doc"):
        doc = _extract_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        doc = _extract_xlsx(file_path)
    elif ext == ".csv":
        doc = _extract_csv(file_path)
    elif ext in (".html", ".htm"):
        doc = _extract_html(file_path)
    elif ext in (".md", ".markdown"):
        doc = _extract_markdown(file_path)
    elif ext in IMAGE_EXTENSIONS:
        doc = await _extract_image(file_path, file_name)
    elif ext in VIDEO_EXTENSIONS:
        doc = await _extract_video(file_path, file_name)
    elif ext in AUDIO_EXTENSIONS:
        doc = await _extract_audio(file_path)
    else:
        try:
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")
            doc = ExtractedDocument(text=text, metadata={"file_type": "text"})
        except Exception as e:
            raise ValueError(f"Unsupported file type: {ext}") from e

    return _sanitize_extracted(doc)
